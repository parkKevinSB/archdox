from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
OUT_DIR = ROOT / "document-engine" / "src" / "main" / "resources" / "templates" / "korean"

FONT = "Malgun Gothic"
TITLE_COLOR = "1F3A5F"
HEADER_FILL = "E8EEF5"
LABEL_FILL = "F2F4F7"
BORDER_COLOR = "5E6A75"
TEXT_COLOR = "111827"
MUTED_COLOR = "4B5563"


def dxa_from_cm(value: float) -> int:
    return round(value / 2.54 * 1440)


def set_east_asia_font(style, font_name: str = FONT) -> None:
    style.font.name = font_name
    rpr = style._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), font_name)
    rfonts.set(qn("w:hAnsi"), font_name)
    rfonts.set(qn("w:eastAsia"), font_name)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.25)
    section.bottom_margin = Cm(1.25)
    section.left_margin = Cm(1.35)
    section.right_margin = Cm(1.35)
    section.header_distance = Cm(0.8)
    section.footer_distance = Cm(0.8)

    normal = doc.styles["Normal"]
    set_east_asia_font(normal)
    normal.font.size = Pt(9.5)
    normal.font.color.rgb = RGBColor.from_string(TEXT_COLOR)
    normal.paragraph_format.space_after = Pt(3)
    normal.paragraph_format.line_spacing = 1.15

    for style_name in ["Title", "Heading 1", "Heading 2", "Heading 3"]:
        set_east_asia_font(doc.styles[style_name])

    doc.styles["Title"].font.size = Pt(16)
    doc.styles["Title"].font.bold = True
    doc.styles["Title"].font.color.rgb = RGBColor.from_string(TITLE_COLOR)
    doc.styles["Title"].paragraph_format.space_after = Pt(6)

    doc.styles["Heading 1"].font.size = Pt(12)
    doc.styles["Heading 1"].font.bold = True
    doc.styles["Heading 1"].font.color.rgb = RGBColor.from_string(TITLE_COLOR)
    doc.styles["Heading 1"].paragraph_format.space_before = Pt(6)
    doc.styles["Heading 1"].paragraph_format.space_after = Pt(4)

    doc.styles["Heading 2"].font.size = Pt(10.5)
    doc.styles["Heading 2"].font.bold = True
    doc.styles["Heading 2"].font.color.rgb = RGBColor.from_string(TITLE_COLOR)
    doc.styles["Heading 2"].paragraph_format.space_before = Pt(5)
    doc.styles["Heading 2"].paragraph_format.space_after = Pt(3)


def set_run_font(run, size: float | None = None, bold: bool = False, color: str = TEXT_COLOR) -> None:
    run.font.name = FONT
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)
    if size is not None:
        run.font.size = Pt(size)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), FONT)
    rfonts.set(qn("w:hAnsi"), FONT)
    rfonts.set(qn("w:eastAsia"), FONT)


def add_para(
    doc: Document,
    text: str,
    style: str | None = None,
    *,
    align: WD_ALIGN_PARAGRAPH | None = None,
    bold: bool = False,
    size: float | None = None,
    color: str = TEXT_COLOR,
):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    return p


def add_title(doc: Document, title: str, form_note: str) -> None:
    add_para(doc, form_note, align=WD_ALIGN_PARAGRAPH.LEFT, size=7.5, color=MUTED_COLOR)
    add_para(doc, title, style="Title", align=WD_ALIGN_PARAGRAPH.CENTER)


def cell_xml(cell, tag: str, attr: str, value: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    element = tc_pr.find(qn(tag))
    if element is None:
        element = OxmlElement(tag)
        tc_pr.append(element)
    element.set(qn(attr), value)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_cm: float) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(dxa_from_cm(width_cm)))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_margins(cell, top: int = 70, bottom: int = 70, start: int = 110, end: int = 110) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in {"top": top, "bottom": bottom, "start": start, "end": end}.items():
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        tag = f"w:{edge}"
        node = borders.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "6")
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), BORDER_COLOR)


def set_table_geometry(table, widths_cm: list[float]) -> None:
    table.autofit = False
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(dxa_from_cm(w) for w in widths_cm)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "0")
    tbl_ind.set(qn("w:type"), "dxa")

    tbl_grid = tbl.tblGrid
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        tbl.insert(0, tbl_grid)
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths_cm:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(dxa_from_cm(width)))
        tbl_grid.append(grid_col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            if index < len(widths_cm):
                set_cell_width(cell, widths_cm[index])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
    set_table_borders(table)


def set_cell_text(
    cell,
    text: str,
    *,
    bold: bool = False,
    fill: str | None = None,
    align: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.LEFT,
    size: float = 8.8,
    color: str = TEXT_COLOR,
) -> None:
    cell.text = ""
    if fill:
        shade_cell(cell, fill)
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_after = Pt(0)
    parts = text.split("\n")
    for idx, part in enumerate(parts):
        if idx:
            p.add_run().add_break()
        run = p.add_run(part)
        set_run_font(run, size=size, bold=bold, color=color)


def add_table(doc: Document, widths_cm: list[float], rows: list[list[dict]]) -> None:
    table = doc.add_table(rows=0, cols=len(widths_cm))
    set_table_geometry(table, widths_cm)
    for row_spec in rows:
        row = table.add_row()
        index = 0
        for spec in row_spec:
            span = spec.get("span", 1)
            cell = row.cells[index]
            if span > 1:
                cell = cell.merge(row.cells[index + span - 1])
            set_cell_text(
                cell,
                spec.get("text", ""),
                bold=spec.get("bold", False),
                fill=spec.get("fill"),
                align=spec.get("align", WD_ALIGN_PARAGRAPH.LEFT),
                size=spec.get("size", 8.8),
                color=spec.get("color", TEXT_COLOR),
            )
            index += span
        for cell_index, cell in enumerate(row.cells):
            if cell_index < len(widths_cm):
                set_cell_width(cell, widths_cm[cell_index])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
    set_table_geometry(table, widths_cm)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def label(text: str, span: int = 1) -> dict:
    return {
        "text": text,
        "bold": True,
        "fill": LABEL_FILL,
        "align": WD_ALIGN_PARAGRAPH.CENTER,
        "span": span,
    }


def value(text: str, span: int = 1, *, align: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.LEFT) -> dict:
    return {"text": text, "span": span, "align": align}


def header(text: str, span: int = 1) -> dict:
    return {
        "text": text,
        "bold": True,
        "fill": HEADER_FILL,
        "align": WD_ALIGN_PARAGRAPH.CENTER,
        "span": span,
    }


def add_section_placeholder(doc: Document, key: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("${" + key + "}")
    set_run_font(run, size=9, bold=False, color="374151")
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)


def add_method_notes(doc: Document, notes: list[str]) -> None:
    add_para(doc, "작성방법", style="Heading 2")
    for note in notes:
        p = doc.add_paragraph(style=None)
        p.paragraph_format.left_indent = Cm(0.25)
        p.paragraph_format.first_line_indent = Cm(-0.25)
        p.paragraph_format.space_after = Pt(1.5)
        run = p.add_run(note)
        set_run_font(run, size=7.8, color=MUTED_COLOR)


def build_construction_daily_log() -> Document:
    doc = Document()
    configure_document(doc)
    add_title(doc, "공 사 감 리 일 지", "건축공사 감리세부기준 [별지 제2호서식]")
    add_table(
        doc,
        [3.0, 5.1, 3.0, 6.9],
        [
            [label("일련번호"), value("${serialNo}"), label("날씨"), value("${weather}")],
            [label("총괄감리책임자"), value("${chiefSupervisorName}  (서명 또는 인)"), label("건축사보"), value("${architectAssistantName}  (서명 또는 인)")],
            [label("공사명"), value("${constructionName}", span=3)],
            [label("공사일자"), value("${inspectionDate} (${inspectionDayOfWeek})"), label("현장"), value("${siteName}")],
        ],
    )
    add_para(doc, "공종 및 세부공정별 감리내용", style="Heading 1")
    add_section_placeholder(doc, "supervisionItemsSection")
    add_table(
        doc,
        [3.2, 14.8],
        [
            [label("특기사항"), value("${specialNotes}")],
            [label("지적사항 및 처리결과"), value("${issueAndAction}")],
        ],
    )
    add_method_notes(
        doc,
        [
            "1. 공종에는 주요공종 및 단위공종 그리고 해당 층수를 기재합니다.",
            "2. 감리항목은 공종별 감리 체크리스트를 기반으로 기재합니다.",
            "3. 감리내용에는 육안검사, 입회, 시험 등 감리내용과 결과를 구체적으로 기재합니다.",
            "4. 특기사항은 특별히 명기되어 있지 아니한 내용의 발생 및 조치사항을 기재합니다.",
            "5. 지적사항 및 처리결과는 재시공, 공사중지 등 지시내용과 처리결과를 기재합니다.",
        ],
    )
    return doc


def build_construction_supervision_report() -> Document:
    doc = Document()
    configure_document(doc)
    add_title(doc, "감 리 보 고 서", "건축공사 감리세부기준 [별지 제1호서식]")
    add_para(
        doc,
        "공사감리보고서는 사용승인신청서와 함께 제출하며, [  ]에는 해당하는 곳에 표시합니다.",
        size=8.2,
        color=MUTED_COLOR,
    )
    add_table(
        doc,
        [3.0, 5.4, 3.0, 6.6],
        [
            [label("신청구분"), value("[  ] 감리중간보고    [  ] 감리완료보고", span=3)],
            [label("허가번호"), value("${permitNumber}"), label("허가일자"), value("${permitDate}")],
            [label("대지위치"), value("${siteAddress}"), label("지번"), value("${lotNumber}")],
            [label("공사명"), value("${constructionName}", span=3)],
        ],
    )
    add_table(
        doc,
        [2.5, 7.0, 4.8, 3.7],
        [
            [header("구분"), header("건축공정"), header("동명칭 및 번호"), header("확인")],
            [value("중간보고\n[  ]", align=WD_ALIGN_PARAGRAPH.CENTER), value("기초공사 철근배근 완료\n[  ] 지붕슬래브 철근배근 완료\n[  ] 층 바닥슬래브 철근배근 완료\n[  ] 거푸집 또는 주춧돌 설치 완료"), value("${buildingName}"), value("${progressType}")],
            [value("완료보고\n[  ]", align=WD_ALIGN_PARAGRAPH.CENTER), value("공사감리기간\n${supervisionStartDate} ~ ${supervisionEndDate}"), value("공사완료일자\n${completionDate}"), value("공사감리자\n${chiefSupervisorName}")],
        ],
    )
    add_table(
        doc,
        [3.4, 14.6],
        [
            [label("관계전문기술자 확인 및 의견"), value("${relationEngineerOpinion}")],
            [label("공사감리자 종합의견"), value("${comprehensiveOpinion}")],
        ],
    )
    add_para(doc, "건축법 제25조 및 같은 법 시행규칙 제19조에 따라 위와 같이 공사감리보고서를 제출합니다.", size=8.5)
    add_table(
        doc,
        [3.2, 5.8, 3.2, 5.8],
        [
            [label("보고일"), value("${reportDate}"), label("보고인"), value("${ownerName}  (서명 또는 인)")],
            [label("공사감리자"), value("${chiefSupervisorName}  (서명 또는 인)", span=3)],
        ],
    )
    doc.add_section(WD_SECTION.NEW_PAGE)
    add_para(doc, "현장 조사 및 법령 기준 확인", style="Title", align=WD_ALIGN_PARAGRAPH.CENTER)
    add_table(
        doc,
        [2.2, 6.3, 4.5, 5.0],
        [
            [header("구분"), header("조사 내용"), header("건축법·조례 기준"), header("완공 후 현황")],
            [value("대지 및 도로", align=WD_ALIGN_PARAGRAPH.CENTER), value("대지의 안전조치 등\n토지굴착 부분에 대한 조치 등\n대지안의 조경 및 건축선"), value("제40조, 제41조, 제46조, 제47조"), value("${fieldInvestigationSummary}")],
            [value("구조내력", align=WD_ALIGN_PARAGRAPH.CENTER), value("철골구조의 품질기준\n주요구조부 및 구조 안전 확인"), value("제48조"), value("${structureStatus}")],
            [value("피난시설", align=WD_ALIGN_PARAGRAPH.CENTER), value("직통계단, 특별피난계단, 출구, 옥상광장, 거실 채광·환기"), value("제49조, 제50조"), value("${evacuationStatus}")],
            [value("방화·내화", align=WD_ALIGN_PARAGRAPH.CENTER), value("방화구획, 내화구조, 방화벽, 마감재료"), value("제49조, 제50조, 제52조"), value("${fireSafetyStatus}")],
            [value("용도·건폐율", align=WD_ALIGN_PARAGRAPH.CENTER), value("용도지역·용도지구, 건폐율, 용적률, 높이제한"), value("국토계획법 제76조 및 건축법 제58조~제61조"), value("${zoningStatus}")],
        ],
    )
    doc.add_section(WD_SECTION.NEW_PAGE)
    add_para(doc, "감리 의견 및 보완 확인", style="Title", align=WD_ALIGN_PARAGRAPH.CENTER)
    add_section_placeholder(doc, "reportOpinionSection")
    add_table(
        doc,
        [3.0, 15.0],
        [
            [label("첨부 및 비고"), value("관계전문기술자 의견서, 현장 사진, 시험성적서, 지시·조치 내역 등 필요한 자료를 첨부합니다.")],
            [label("특기사항"), value("${specialNotes}")],
        ],
    )
    return doc


def build_demolition_safety_checklist() -> Document:
    doc = Document()
    configure_document(doc)
    add_title(doc, "해체공사 안전점검표", "건축물 해체공사 감리업무 [별지 제1호서식]")
    add_table(
        doc,
        [3.0, 5.5, 3.0, 6.5],
        [
            [label("점검일자"), value("${safetyInspectionDate}"), label("점검위치"), value("${inspectionLocation}")],
            [label("감리자"), value("${supervisorName}  (서명)"), label("해체작업자"), value("${demolitionWorkerName}  (서명)")],
            [label("작업단계"), value("${safetyCheckStage}"), label("조치요약"), value("${correctiveAction}")],
        ],
    )
    add_section_placeholder(doc, "safetyChecklistSection")
    add_method_notes(
        doc,
        [
            "1. 안전점검표에는 하부보강 잭서포트 재원 및 설치 간격, 적용 층수, 해체장비 이동구간 잔재물 적재 높이와 하중, 보강 상세도면을 포함합니다.",
            "2. 세부 검사항목은 해체작업순서에 따른 주요사항과 잔재물 허용범위를 기재합니다.",
            "3. 조치사항은 부적합사항에 대한 작업요청 사항을 기입하고 수정·보완사항을 표시합니다.",
            "※ 현장여건에 따라 필수확인점 변경이 필요한 경우 해체작업자 및 관리자와 협의하여 변경할 수 있습니다.",
        ],
    )
    return doc


def build_demolition_daily_log() -> Document:
    doc = Document()
    configure_document(doc)
    add_title(doc, "공 사 감 리 일 지", "건축물 해체공사 감리업무 [별지 제2호서식]")
    add_table(
        doc,
        [3.0, 5.5, 3.0, 6.5],
        [
            [label("공사감리자"), value("${supervisorName}  (서명 또는 인)"), label("감리원"), value("${inspectorName}  (서명 또는 인)")],
            [label("공사명"), value("${constructionName}", span=3)],
            [label("공사일자"), value("${inspectionDate} (${inspectionDayOfWeek})"), label("날씨"), value("${weather}")],
            [label("작업사항"), value("${workDescription}", span=3)],
        ],
    )
    add_para(doc, "공종별 감리착안사항 및 감리내용", style="Heading 1")
    add_section_placeholder(doc, "supervisionItemsSection")
    add_table(
        doc,
        [3.2, 14.8],
        [
            [label("특기사항"), value("${specialNotes}")],
            [label("지적사항 및 처리결과"), value("${issueAndAction}")],
        ],
    )
    add_method_notes(
        doc,
        [
            "1. 공종에는 주요공종 및 단위공종을 기재합니다.",
            "2. 감리착안사항은 공사감리의 주안점 및 점검계획을 기재합니다.",
            "3. 특기사항은 특별히 명기되어 있지 아니한 내용의 발생 및 조치사항을 기재합니다.",
            "4. 지적사항 및 처리결과는 재시공, 공사중지 등 지시내용과 처리결과를 기재합니다.",
            "※ 필수확인점에 해당하는 경우에는 반드시 작성하여야 합니다.",
        ],
    )
    return doc


def build_demolition_completion_report() -> Document:
    doc = Document()
    configure_document(doc)
    add_title(doc, "건축물 해체감리완료 보고서", "건축물 해체공사 감리업무 [별지 제3호서식]")
    add_table(
        doc,
        [3.2, 5.8, 3.2, 5.8],
        [
            [header("감리자", span=4)],
            [label("성명(대표자명)"), value("${supervisorName}"), label("상호명/자격번호"), value("${supervisorOfficeName} / ${supervisorLicenseNumber}")],
            [label("주소"), value("${supervisorAddress}", span=3)],
            [label("전화번호"), value("${supervisorPhone}"), label("신고번호"), value("${permitNumber}")],
            [header("공사시공자", span=4)],
            [label("성명(대표자명)"), value("${contractorName}"), label("상호명/면허번호"), value("${contractorOfficeName} / ${contractorLicenseNumber}")],
            [label("주소"), value("${contractorAddress}", span=3)],
            [label("전화번호"), value("${contractorPhone}", span=3)],
        ],
    )
    add_table(
        doc,
        [3.2, 5.8, 3.2, 5.8],
        [
            [header("공사감리 용역현황", span=4)],
            [label("용역명"), value("${serviceName}"), label("현장주소"), value("${siteAddress}")],
            [label("용역개요"), value("${serviceOverview}", span=3)],
            [label("공사기간"), value("${constructionStartDate} ~ ${constructionEndDate}"), label("공사금액"), value("${constructionAmount} 천원")],
            [label("감리기간"), value("${supervisionStartDate} ~ ${supervisionEndDate}"), label("감리금액"), value("${supervisionAmount} 천원")],
        ],
    )
    add_para(doc, "감리원 배치현황", style="Heading 1")
    add_section_placeholder(doc, "supervisorDeploymentSection")
    add_table(
        doc,
        [3.0, 15.0],
        [
            [label("종합의견"), value("${comprehensiveOpinion}")],
            [label("제출일"), value("${reportDate}")],
        ],
    )
    add_para(doc, "건축물관리법 제32조제5항에 따라 위와 같이 건축물 해체감리완료보고서를 제출합니다.", size=8.5)
    add_method_notes(
        doc,
        [
            "첨부: 1. 해체공사 및 감리수행 결과",
            "첨부: 2. 안전점검표",
            "첨부: 3. 감리업무일지",
            "첨부: 4. 각종 반입자재 규격 및 반입장비 제원",
            "첨부: 5. 공사 현황 사진 및 동영상",
            "첨부: 6. 기타 감리자 의견서",
        ],
    )
    return doc


TEMPLATES = [
    (
        "korean-construction-supervision-report-appendix-1.docx",
        build_construction_supervision_report,
    ),
    (
        "korean-construction-daily-supervision-log-appendix-2.docx",
        build_construction_daily_log,
    ),
    (
        "korean-demolition-safety-checklist-appendix-1.docx",
        build_demolition_safety_checklist,
    ),
    (
        "korean-demolition-daily-supervision-log-appendix-2.docx",
        build_demolition_daily_log,
    ),
    (
        "korean-demolition-completion-report-appendix-3.docx",
        build_demolition_completion_report,
    ),
]


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    for filename, builder in TEMPLATES:
        document = builder()
        document.save(OUT_DIR / filename)
        print(OUT_DIR / filename)


if __name__ == "__main__":
    main()
