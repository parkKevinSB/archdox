from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
OUT_DIR = ROOT / "document-engine" / "src" / "main" / "resources" / "templates" / "korean"

FONT = "Malgun Gothic"
TEXT = "111827"
MUTED = "4B5563"
TITLE = "1F3A5F"
HEADER_FILL = "E8EEF5"
LABEL_FILL = "F2F4F7"
ACCENT_FILL = "EFE7C6"
BORDER = "5E6A75"


def dxa_from_cm(value: float) -> int:
    return round(value / 2.54 * 1440)


def set_run_font(run, size: float | None = None, bold: bool = False, color: str = TEXT) -> None:
    run.font.name = FONT
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)
    if size is not None:
        run.font.size = Pt(size)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("ascii", "hAnsi", "eastAsia"):
        rfonts.set(qn(f"w:{attr}"), FONT)


def set_style_font(style, size: float, color: str = TEXT, bold: bool = False) -> None:
    style.font.name = FONT
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = RGBColor.from_string(color)
    rpr = style._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("ascii", "hAnsi", "eastAsia"):
        rfonts.set(qn(f"w:{attr}"), FONT)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.2)
    section.bottom_margin = Cm(1.2)
    section.left_margin = Cm(1.35)
    section.right_margin = Cm(1.35)
    section.header_distance = Cm(0.7)
    section.footer_distance = Cm(0.7)

    normal = doc.styles["Normal"]
    set_style_font(normal, 9.2)
    normal.paragraph_format.space_after = Pt(3)
    normal.paragraph_format.line_spacing = 1.12

    set_style_font(doc.styles["Title"], 16, TITLE, True)
    doc.styles["Title"].paragraph_format.space_after = Pt(5)
    set_style_font(doc.styles["Heading 1"], 11.5, TITLE, True)
    doc.styles["Heading 1"].paragraph_format.space_before = Pt(6)
    doc.styles["Heading 1"].paragraph_format.space_after = Pt(4)
    set_style_font(doc.styles["Heading 2"], 10.2, TITLE, True)
    doc.styles["Heading 2"].paragraph_format.space_before = Pt(5)
    doc.styles["Heading 2"].paragraph_format.space_after = Pt(3)


def paragraph(
    doc: Document,
    text: str,
    style: str | None = None,
    *,
    align: WD_ALIGN_PARAGRAPH | None = None,
    size: float | None = None,
    bold: bool = False,
    color: str = TEXT,
):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    return p


def title(doc: Document, form_note: str, title_text: str) -> None:
    paragraph(doc, form_note, align=WD_ALIGN_PARAGRAPH.LEFT, size=7.5, color=MUTED)
    paragraph(doc, title_text, style="Title", align=WD_ALIGN_PARAGRAPH.CENTER)


def set_cell_margins(cell, top: int = 70, bottom: int = 70, start: int = 110, end: int = 110) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in {"top": top, "bottom": bottom, "start": start, "end": end}.items():
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "6")
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), BORDER)


def set_table_geometry(table, widths_cm: list[float]) -> None:
    table.autofit = False
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(dxa_from_cm(w) for w in widths_cm)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_grid = tbl.tblGrid
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        tbl.insert(0, tbl_grid)
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths_cm:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(dxa_from_cm(width)))
        tbl_grid.append(grid_col)

    set_table_borders(table)


def set_cell_width(cell, width_cm: float) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(dxa_from_cm(width_cm)))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_text(
    cell,
    text: str,
    *,
    bold: bool = False,
    fill: str | None = None,
    align: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.LEFT,
    size: float = 8.7,
    color: str = TEXT,
) -> None:
    cell.text = ""
    if fill:
        shade_cell(cell, fill)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(cell)
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_after = Pt(0)
    for index, line in enumerate(text.split("\n")):
        if index:
            p.add_run().add_break()
        run = p.add_run(line)
        set_run_font(run, size=size, bold=bold, color=color)


def label(text: str, span: int = 1) -> dict:
    return {"text": text, "bold": True, "fill": LABEL_FILL, "align": WD_ALIGN_PARAGRAPH.CENTER, "span": span}


def header(text: str, span: int = 1) -> dict:
    return {"text": text, "bold": True, "fill": HEADER_FILL, "align": WD_ALIGN_PARAGRAPH.CENTER, "span": span}


def value(text: str, span: int = 1, *, align: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.LEFT) -> dict:
    return {"text": text, "span": span, "align": align}


def table(doc: Document, widths_cm: list[float], rows: list[list[dict]]) -> None:
    t = doc.add_table(rows=0, cols=len(widths_cm))
    set_table_geometry(t, widths_cm)
    for row_spec in rows:
        row = t.add_row()
        index = 0
        for spec in row_spec:
            span = spec.get("span", 1)
            cell = row.cells[index]
            if span > 1:
                cell = cell.merge(row.cells[index + span - 1])
            set_cell_text(
                cell,
                spec.get("text", ""),
                bold=spec.get("bold", False),
                fill=spec.get("fill"),
                align=spec.get("align", WD_ALIGN_PARAGRAPH.LEFT),
                size=spec.get("size", 8.7),
                color=spec.get("color", TEXT),
            )
            index += span
        for i, cell in enumerate(row.cells):
            if i < len(widths_cm):
                set_cell_width(cell, widths_cm[i])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
    set_table_geometry(t, widths_cm)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def placeholder(doc: Document, key: str, *, fill: str | None = None) -> None:
    t = doc.add_table(rows=1, cols=1)
    set_table_geometry(t, [18.0])
    cell = t.rows[0].cells[0]
    set_cell_text(
        cell,
        "${" + key + "}",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        size=8.8,
        color="374151",
        fill=fill,
    )
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def method_notes(doc: Document, notes: list[str]) -> None:
    paragraph(doc, "작성방법", style="Heading 2")
    for note in notes:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.28)
        p.paragraph_format.first_line_indent = Cm(-0.28)
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run(note)
        set_run_font(run, size=7.6, color=MUTED)


def build_construction_daily_log() -> Document:
    doc = Document()
    configure_document(doc)
    title(doc, "건축공사 감리세부기준 [별지 제2호서식]", "공사감리일지")
    table(
        doc,
        [3.0, 5.1, 3.0, 6.9],
        [
            [label("일련번호"), value("${serialNo}"), label("날씨"), value("${weather}")],
            [label("총괄감리책임자"), value("${chiefSupervisorName}  (서명 또는 인)"), label("건축사보"), value("${architectAssistantName}  (서명 또는 인)")],
            [label("공사명"), value("${constructionName}", span=3)],
            [label("공사일자"), value("${inspectionDate} (${inspectionDayOfWeek})"), label("현장"), value("${siteName}")],
        ],
    )
    paragraph(doc, "공종 및 세부공정별 감리내용", style="Heading 1")
    placeholder(doc, "supervisionItemsSection")
    table(
        doc,
        [3.2, 14.8],
        [
            [label("특기사항"), value("${specialNotes}")],
            [label("지적사항 및 처리결과"), value("${issueAndAction}")],
        ],
    )
    paragraph(doc, "체크 항목별 사진", style="Heading 1")
    placeholder(doc, "checklistPhotoSection", fill=ACCENT_FILL)
    paragraph(doc, "현장 사진", style="Heading 1")
    placeholder(doc, "photoSection")
    method_notes(
        doc,
        [
            "1. 공종에는 주요공종 및 단위공종, 해당 층수 또는 구역을 기록합니다.",
            "2. 감리항목은 공종별 감리 체크리스트와 현장 확인사항을 기준으로 작성합니다.",
            "3. 감리내용에는 도면ㆍ시방ㆍ자재ㆍ시공상태 확인 내용과 조치 결과를 구체적으로 기록합니다.",
            "4. 특기사항은 별도로 명기하여야 하는 내용, 현장 여건, 협의사항을 기록합니다.",
            "5. 지적사항 및 처리결과에는 시정 지시, 보완 요청, 처리 완료 여부를 함께 기록합니다.",
        ],
    )
    return doc


def build_construction_supervision_report() -> Document:
    doc = Document()
    configure_document(doc)
    title(doc, "건축공사 감리세부기준 [별지 제1호서식]", "감리보고서")
    paragraph(doc, "공사감리보고서는 사용승인 신청 시 함께 제출하며, 해당되는 보고 구분에 표시합니다.", size=8.2, color=MUTED)
    table(
        doc,
        [3.0, 5.4, 3.0, 6.6],
        [
            [label("신청구분"), value("[  ] 감리중간보고    [  ] 감리완료보고", span=3)],
            [label("허가번호"), value("${permitNumber}"), label("허가일자"), value("${permitDate}")],
            [label("대지위치"), value("${siteAddress}"), label("지번"), value("${lotNumber}")],
            [label("공사명"), value("${constructionName}", span=3)],
        ],
    )
    table(
        doc,
        [2.5, 7.0, 4.8, 3.7],
        [
            [header("구분"), header("건축공정"), header("동명칭 및 번호"), header("확인")],
            [value("중간보고\n[  ]", align=WD_ALIGN_PARAGRAPH.CENTER), value("기초공사 철근배근 완료\n[  ] 지붕슬래브 철근배근 완료\n[  ] 주요 구조부 공정 완료\n[  ] 거푸집 또는 주심재 설치 완료"), value("${constructionName}"), value("확인")],
            [value("완료보고\n[  ]", align=WD_ALIGN_PARAGRAPH.CENTER), value("공사감리기간\n${supervisionStartDate} ~ ${supervisionEndDate}"), value("감리자\n${chiefSupervisorName}"), value("${supervisorName}")],
        ],
    )
    table(
        doc,
        [3.4, 14.6],
        [
            [label("관계전문기술자 확인 및 의견"), value("${relationEngineerOpinion}")],
            [label("공사감리자 종합의견"), value("${comprehensiveOpinion}")],
        ],
    )
    paragraph(doc, "건축법 및 관계 법령에 따라 위와 같이 공사감리보고서를 제출합니다.", size=8.5)
    table(
        doc,
        [3.2, 5.8, 3.2, 5.8],
        [
            [label("보고일"), value("${reportDate}"), label("건축주"), value("                    (서명 또는 인)")],
            [label("공사감리자"), value("${chiefSupervisorName}  (서명 또는 인)", span=3)],
        ],
    )
    doc.add_page_break()
    paragraph(doc, "현장 조사 및 법령 기준 확인", style="Title", align=WD_ALIGN_PARAGRAPH.CENTER)
    table(
        doc,
        [2.2, 6.3, 4.5, 5.0],
        [
            [header("구분"), header("조사 내용"), header("관련 기준"), header("시공 및 확인 현황")],
            [value("대지 및 도로", align=WD_ALIGN_PARAGRAPH.CENTER), value("대지의 안전조치\n도로ㆍ인접지와의 관계\n대지 내 조경 및 배치"), value("건축법령 및 허가조건"), value("감리 체크 결과 참조")],
            [value("구조", align=WD_ALIGN_PARAGRAPH.CENTER), value("주요 구조부 시공상태\n철근ㆍ콘크리트ㆍ철골 등 구조 안전 확인"), value("구조도서 및 감리 체크 기준"), value("감리 체크 결과 참조")],
            [value("피난ㆍ방화", align=WD_ALIGN_PARAGRAPH.CENTER), value("피난계단, 출구, 방화구획, 방화문 등"), value("건축법령 피난ㆍ방화 기준"), value("감리 체크 결과 참조")],
            [value("설비ㆍ마감", align=WD_ALIGN_PARAGRAPH.CENTER), value("주요 설비, 마감, 사용승인 전 보완 필요사항"), value("설계도서 및 관계 법령"), value("${specialNotes}")],
        ],
    )
    doc.add_page_break()
    paragraph(doc, "감리 의견 및 보완 확인", style="Title", align=WD_ALIGN_PARAGRAPH.CENTER)
    placeholder(doc, "reportOpinionSection")
    paragraph(doc, "첨부 사진", style="Heading 1")
    placeholder(doc, "photoSection")
    table(
        doc,
        [3.0, 15.0],
        [
            [label("첨부 및 비고"), value("관계전문기술자 의견서, 현장 사진, 시험성적서, 지시공문 등 필요한 자료를 첨부합니다.")],
            [label("특기사항"), value("${specialNotes}")],
        ],
    )
    return doc


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    outputs = {
        "korean-construction-daily-supervision-log-appendix-2.docx": build_construction_daily_log,
        "korean-construction-supervision-report-appendix-1.docx": build_construction_supervision_report,
    }
    for filename, builder in outputs.items():
        path = OUT_DIR / filename
        builder().save(path)
        print(path)


if __name__ == "__main__":
    main()
